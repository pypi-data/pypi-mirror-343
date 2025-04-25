import docx
import pytest
import numpy as np
from unittest.mock import patch, MagicMock
from vectoriz.files import FileArgument
from vectoriz.files import FilesFeature


class TestFileArgument:
    def test_add_data_appends_to_lists(self):
        file_arg = FileArgument()
        filename = "test.txt"
        text = "Test content"

        with patch.object(
            FileArgument, "_create_embedding", return_value=np.array([0.1, 0.2, 0.3])
        ):
            file_arg.add_data(filename, text)

            assert file_arg.chunk_names == [filename]
            assert file_arg.text_list == [text]
            assert len(file_arg.embeddings) == 1
            np.testing.assert_array_equal(
                file_arg.embeddings[0], np.array([0.1, 0.2, 0.3])
            )

    def test_add_data_multiple_entries(self):
        file_arg = FileArgument(
            ["existing.txt"], ["existing content"], [np.array([0.5, 0.5, 0.5])]
        )
        filename = "new.txt"
        text = "New content"

        with patch.object(
            FileArgument, "_create_embedding", return_value=np.array([0.7, 0.8, 0.9])
        ):
            file_arg.add_data(filename, text)
            assert file_arg.chunk_names == ["existing.txt", "new.txt"]
            assert file_arg.text_list == ["existing content", "New content"]
            assert len(file_arg.embeddings) == 2
            np.testing.assert_array_equal(
                file_arg.embeddings[1], np.array([0.7, 0.8, 0.9])
            )

    def test_add_data_calls_create_embedding(self):
        file_arg = FileArgument()
        filename = "test.txt"
        text = "Test content"

        with patch.object(FileArgument, "_create_embedding") as mock_create_embedding:
            mock_create_embedding.return_value = np.array([0.1, 0.2, 0.3])
            file_arg.add_data(filename, text)
            mock_create_embedding.assert_called_once_with(text)

    def test_create_embedding_returns_numpy_array(self):
        file_arg = FileArgument()
        text = "Test content"

        with patch("vectoriz.files.TokenTransformer") as mock_transformer:
            mock_instance = mock_transformer.return_value
            mock_instance.text_to_embeddings.return_value = [np.array([0.1, 0.2, 0.3])]

            result = file_arg._create_embedding(text)

            assert isinstance(result, np.ndarray)
            np.testing.assert_array_equal(result, np.array([0.1, 0.2, 0.3]))
            mock_instance.text_to_embeddings.assert_called_once_with([text])

    def test_create_embedding_handles_empty_text(self):
        file_arg = FileArgument()
        text = ""

        with patch("vectoriz.files.TokenTransformer") as mock_transformer:
            mock_instance = mock_transformer.return_value
            mock_instance.text_to_embeddings.return_value = [np.array([0.0, 0.0, 0.0])]

            result = file_arg._create_embedding(text)

            assert isinstance(result, np.ndarray)
            mock_instance.text_to_embeddings.assert_called_once_with([""])

    def test_create_embedding_instantiates_token_transformer(self):
        file_arg = FileArgument()
        text = "Test content"

        with patch("vectoriz.files.TokenTransformer") as mock_transformer:
            mock_instance = mock_transformer.return_value
            mock_instance.text_to_embeddings.return_value = [np.array([0.1, 0.2, 0.3])]

            file_arg._create_embedding(text)

            mock_transformer.assert_called_once()


class TestFilesFeature:
    def test_extract_txt_content_reads_file_correctly(self, tmp_path):
        test_content = "This is test content"
        test_file = tmp_path / "test.txt"
        test_file.write_text(test_content)
        files_feature = FilesFeature()
        result = files_feature._extract_txt_content(test_file)
        assert result == {"file": "test.txt", "content": test_content}

    def test_extract_txt_content_with_unicode_chars(self, tmp_path):
        test_content = "Unicode content: àáâãäåæç"
        test_file = tmp_path / "unicode.txt"
        test_file.write_text(test_content, encoding="utf-8")
        files_feature = FilesFeature()
        result = files_feature._extract_txt_content(test_file)
        assert result == {"file": "unicode.txt", "content": test_content}

    def test_extract_txt_content_raises_file_not_found(self):
        files_feature = FilesFeature()
        with pytest.raises(FileNotFoundError):
            files_feature._extract_txt_content(
                "/non_existent_dir/non_existent_file.txt"
            )

    def test_extract_docx_content_reads_file_correctly(self, tmp_path, monkeypatch):
        mock_doc = MagicMock()
        mock_paragraph1 = MagicMock()
        mock_paragraph1.text = "Paragraph 1"
        mock_paragraph2 = MagicMock()
        mock_paragraph2.text = "Paragraph 2"
        mock_doc.paragraphs = [mock_paragraph1, mock_paragraph2]

        monkeypatch.setattr(docx, "Document", lambda _: mock_doc)
        files_feature = FilesFeature()
        path = tmp_path / "test.docx"
        result = files_feature._extract_docx_content(path)

        assert result == {"file": "test.docx", "content": "Paragraph 1\nParagraph 2"}

    def test_extract_docx_content_skips_empty_paragraphs(self, tmp_path, monkeypatch):
        mock_doc = MagicMock()
        mock_paragraph1 = MagicMock()
        mock_paragraph1.text = "Paragraph 1"
        mock_paragraph2 = MagicMock()
        mock_paragraph2.text = "   "
        mock_paragraph3 = MagicMock()
        mock_paragraph3.text = "Paragraph 3"
        mock_doc.paragraphs = [mock_paragraph1, mock_paragraph2, mock_paragraph3]

        monkeypatch.setattr(docx, "Document", lambda _: mock_doc)
        files_feature = FilesFeature()
        path = tmp_path / "test.docx"
        result = files_feature._extract_docx_content(path)

        assert result == {"file": "test.docx", "content": "Paragraph 1\nParagraph 3"}

    def test_extract_docx_content_exception_handling(self, tmp_path, monkeypatch):
        def mock_document(_):
            raise Exception("Failed to open document")

        monkeypatch.setattr(docx, "Document", mock_document)

        files_feature = FilesFeature()
        with pytest.raises(Exception):
            path = tmp_path / "/invalid.docx"
            files_feature._extract_docx_content(path)

    def test_extract_docx_content_with_no_paragraphs(self, tmp_path, monkeypatch):
        mock_doc = MagicMock()
        mock_doc.paragraphs = []
        monkeypatch.setattr(docx, "Document", lambda _: mock_doc)
        files_feature = FilesFeature()
        path = tmp_path / "empty.docx"
        result = files_feature._extract_docx_content(path)
        assert result == {"file": "empty.docx", "content": ""}

    def test_extract_markdown_content_reads_file_correctly(self, tmp_path):
        test_content = "# Markdown Title\nThis is some markdown content."
        test_file = tmp_path / "test.md"
        test_file.write_text(test_content)
        files_feature = FilesFeature()
        path = tmp_path / "test.md"
        result = files_feature._extract_markdown_content(path)
        assert result == {"file": "test.md", "content": test_content}

    def test_extract_markdown_content_with_unicode_chars(self, tmp_path):
        test_content = "# Unicode Title\nContent with unicode: àáâãäåæç"
        test_file = tmp_path / "unicode.md"
        test_file.write_text(test_content, encoding="utf-8")
        files_feature = FilesFeature()
        path = tmp_path / "unicode.md"
        result = files_feature._extract_markdown_content(path)
        assert result == {"file": "unicode.md", "content": test_content}

    def test_extract_markdown_content_raises_file_not_found(self):
        files_feature = FilesFeature()
        with pytest.raises(FileNotFoundError):
            path = str("/non_existent_dir/non_existent_file.md")
            files_feature._extract_markdown_content(path)

    def test_extract_markdown_content_handles_empty_file(self, tmp_path):
        test_file = tmp_path / "empty.md"
        test_file.write_text("")
        files_feature = FilesFeature()
        result = files_feature._extract_markdown_content(test_file)
        assert result == {'file': 'empty.md', 'content': ''}

    def test_extract_markdown_content_raises_unicode_decode_error(self, tmp_path):
        test_file = tmp_path / "invalid_encoding.md"
        test_file.write_bytes(b"\x80\x81\x82")
        files_feature = FilesFeature()
        with pytest.raises(UnicodeDecodeError):
            path = str(tmp_path / "invalid_encoding.md")
            files_feature._extract_markdown_content(path)

    def test_load_markdown_files_from_directory_loads_files_correctly(self, tmp_path):
        test_content_1 = "# Title 1\nContent 1"
        test_content_2 = "# Title 2\nContent 2"
        test_file_1 = tmp_path / "file1.md"
        test_file_2 = tmp_path / "file2.md"
        test_file_1.write_text(test_content_1)
        test_file_2.write_text(test_content_2)

        files_feature = FilesFeature()
        result = files_feature.load_markdown_files_from_directory(tmp_path)

        assert len(result.chunk_names) == 2
        assert len(result.text_list) == 2
        assert test_file_1.name in result.chunk_names
        assert test_file_2.name in result.chunk_names
        assert test_content_1 in result.text_list
        assert test_content_2 in result.text_list

    def test_load_markdown_files_from_directory_skips_non_markdown_files(self, tmp_path):
        test_content_md = "# Markdown Content"
        test_content_txt = "Text Content"
        test_file_md = tmp_path / "file.md"
        test_file_txt = tmp_path / "file.txt"
        test_file_md.write_text(test_content_md)
        test_file_txt.write_text(test_content_txt)

        files_feature = FilesFeature()
        result = files_feature.load_markdown_files_from_directory(tmp_path)

        assert len(result.chunk_names) == 1
        assert len(result.text_list) == 1
        assert test_file_md.name in result.chunk_names
        assert test_content_md in result.text_list
        assert test_file_txt.name not in result.chunk_names

    def test_load_markdown_files_from_directory_handles_empty_directory(self, tmp_path):
        files_feature = FilesFeature()
        result = files_feature.load_markdown_files_from_directory(tmp_path)

        assert len(result.chunk_names) == 0
        assert len(result.text_list) == 0

    def test_load_markdown_files_from_directory_handles_empty_markdown_file(self, tmp_path):
        test_file = tmp_path / "empty.md"
        test_file.write_text("")

        files_feature = FilesFeature()
        result = files_feature.load_markdown_files_from_directory(str(tmp_path))

        assert len(result.chunk_names) == 1
        assert len(result.text_list) == 1
        assert test_file.name in result.chunk_names
        assert result.text_list[0] == ""

    def test_load_markdown_files_from_directory_with_verbose_output(self, tmp_path, capsys):
        test_content = "# Markdown Content"
        test_file = tmp_path / "file.md"
        test_file.write_text(test_content)

        files_feature = FilesFeature()
        files_feature.load_markdown_files_from_directory(str(tmp_path), verbose=True)

        captured = capsys.readouterr()
        assert "Loaded Markdown file: file.md" in captured.out


