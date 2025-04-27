import json
import os
import time
from os import path

import dill
import pyperclip
from docx import Document
from docx.shared import Cm
from lxml import etree

header_comment = '#'


def dillWrite(file, obj, base_dir='.'):
    if not os.path.exists(base_dir):
        os.makedirs(base_dir, exist_ok=True)
    if type(file) == tuple or type(file) == list:
        for f, o in zip(file, obj):
            full_path = f'{base_dir}/{f}' if f.endswith('.pkl') else f'{base_dir}/{f}.pkl'
            with open(full_path, 'wb') as ff:
                ff.write(dill.dumps(o))
    else:
        full_path = f'{base_dir}/{file}' if file.endswith('.pkl') else f'{base_dir}/{file}.pkl'
        with open(full_path, 'wb') as ff:
            ff.write(dill.dumps(obj))


def dillRead(file, base_dir='.'):
    if type(file) == tuple or type(file) == list:
        obj = []
        for f in file:
            full_path = f'{base_dir}/{f}' if f.endswith('.pkl') else f'{base_dir}/{f}.pkl'
            with open(full_path, 'rb') as ff:
                obj.append(dill.loads(ff.read()))

        return tuple(obj)
    else:
        full_path = f'{base_dir}/{file}' if file.endswith('.pkl') else f'{base_dir}/{file}.pkl'
        with open(full_path, 'rb') as ff:
            obj = dill.loads(ff.read())
        return obj


def showProgress(total_seconds):
    """
    show a dynamic progress bar with percentage in command line interface
    :param total_seconds: the whole time needed in seconds
    :return:
    """
    for i in range(100):
        print(chr(9600) * i, i, '%', '-' * (100 - i), end='')
        time.sleep(total_seconds / 100)
        print('\r', end='\r')
    else:
        print(chr(9600) * 100, 100, '%')


def printing(*args):
    print(*args)
    print('变量类型是：{}'.format(type(*args)))
    print('内存地址是：{}'.format(id(*args)))
    try:
        print('长度是:{}'.format(len(*args)))
    except:
        print('没有长度')
    try:
        print('变量的所有属性和方法名：{}'.format(*args.__dir__()))
    except:
        print('没有属性和方法')


def color_print(*args, mode: str = 'default', font_color: str = 'white', background_color: str = 'black', sep=' ',
                end='\n'):
    """
    to print values in a more colorful and rich fathon
    :param args: the values to print
    :param mode: the chosen mode, like default, contrast, underline, italic, highlight, blink, strike and their negativity
    :param font_color: the chosen font color, like black, red, green, blue, yellow, magenta, cyan, grey, white and transparent
    :param background_color: the chosen background color, like black, red, green, blue, yellow, magenta, cyan, grey and box
    :param sep: the seperator string
    :param end: the end string
    :return: None
    """
    # mapping the integer of parameters to string for reason of easy memorization
    mode_map = {'default': 0,
                'highlight': 1,
                'non_bold': 22,
                'italic': 3,
                'underline': 4,
                'blink': 5,
                'non_blink': 25,
                'contrast': 7,
                'strike': 9,
                'non_contrast': 27
                }
    font_color_map = {'black': 30,
                      'red': 31,
                      'green': 32,
                      'yellow': 33,
                      'blue': 34,
                      'magenta': 35,
                      'cyan': 36,
                      'grey': 37,
                      'transparent': 38,
                      'white': 39}
    background_color_map = {'black': 40,
                            'red': 41,
                            'green': 42,
                            'yellow': 43,
                            'blue': 44,
                            'magenta': 45,
                            'cyan': 46,
                            'grey': 47,
                            'box': 51}
    for arg in args:
        # if not the last value to print, use the sep value as end
        if arg is not args[-1]:
            formatted_value = f'\033[{mode_map[mode]};{font_color_map[font_color]};{background_color_map[background_color]}m{arg}\033[0m'
            print(formatted_value, sep='', end=sep)
        # if the last value to print, use the end value as end
        else:
            formatted_value = f'\033[{mode_map[mode]};{font_color_map[font_color]};{background_color_map[background_color]}m{arg}\033[0m'
            print(formatted_value, sep='', end=end)


def findFilesRecurisive(path: str):
    """
    find all the files in the path
    :param path: a folder name
    :return:
    """
    all_files = []

    def findAllFiles(path):
        nonlocal all_files
        dirs = os.listdir(path)
        for file in dirs:
            try:
                full_path = os.path.join(path, file)
                if os.path.isfile(full_path):
                    all_files.append(full_path)
                elif os.path.isdir(full_path):
                    findAllFiles(full_path)
            except:
                continue
        return all_files

    return findAllFiles(path)


def findFiles(path: str) -> list:
    gen = os.walk(path)  # 生成器，包含的是元组，没有完整目录。元组第一个是目录（字符串），第二个是包含的子文件夹（列表），第三个是文件列表
    files = []
    for i in gen:
        if i[2]:  # 表示有文件
            for file in i[2]:
                if not file.startswith(('~', '.')):
                    files.append(f'{i[0]}/{file}')  # 完整目录名加上文件名
    return files


def findAllTypeFiles(path: str, extensions) -> list:
    """
    get the list of given type files in the folder path
    :param path: folder name
    :param extensions: file type string or tuple of strings [lowercase]
    :return:
    """
    all_files = findFiles(path)
    if extensions is not None:
        all_type_files = []
        for file in all_files:
            if file.endswith(extensions):
                all_type_files.append(file)

        return all_type_files
    return all_files


def printx(*args, sep=' ', end='\n'):
    """
    :param  args:
    :param sep:
    :param end:
    :return
     """
    globs = globals().items()
    for arg in args:
        for k, v in globs:
            if v == arg and k != '_':
                print(f'{k} value is: {v}', sep=sep, end=end)


def openDocx(file):  # 打开指定DOCX文件，读取文字
    text = ''
    try:
        doc = Document(file)
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += ('    ' + paragraph.text + '\n')

        if doc.tables:
            for tab in doc.tables:
                rows_num = len(tab.rows)
                columns_num = len(tab.columns)
                cell_set = set()
                for row in range(rows_num):
                    for column in range(columns_num):
                        cell_set.add(tab.cell(row, column))
                for i in cell_set:
                    text += (i.text + '\t')

        if not text:
            if doc._body._element.xml:
                body_xml_str = doc._body._element.xml  # 获取body中的xml
                body_xml = etree.fromstring(body_xml_str)  # 转换成lxml结点
                text = etree.tounicode(body_xml)  # 转换成字符串格式
    except:
        text = ''
    return text  # 返回字符串


def searchTextInFile(keyword: str, file) -> int:
    try:
        with open(file, 'r', encoding='utf8') as f:
            content = f.read()
        return content.find(keyword)
    except:
        return -1


def searchTextInFiles(keyword: str, file_list: list) -> list:
    results = []
    for file in file_list:
        ind = searchTextInFile(keyword, file)
        if ind > 0:
            results.append((file, ind))
    return results


def getClipboard():
    """获取剪贴板数据"""
    data = pyperclip.paste()  # 主要这里差别
    return data


def listenClipboard():
    """后台脚本：每隔0.2秒，读取剪切板文本，检查有无指定字符或字符串，如果有则执行替换"""
    # recent_txt 存放最近一次剪切板文本，初始化值只多执行一次paste函数读取和替换
    recent_txt = pyperclip.paste()
    while True:
        # txt 存放当前剪切板文本
        txt = pyperclip.paste()
        # 剪切板内容和上一次对比如有变动，再进行内容判断，判断后如果发现有指定字符在其中的话，再执行替换
        if txt != recent_txt:
            # print(f'txt:{txt}')
            recent_txt = txt  # 没查到要替换的子串，返回None
            return recent_txt

        # 检测间隔（延迟0.2秒）
        time.sleep(0.2)


def readTxt(txt):
    """all the files that can be opened in a txt way, like ini, conf, qss, css, csv
    :param txt:
    :return:
    """
    with open(txt, 'r') as f:
        return f.read()


def countOccurrence(word_list, content):
    """
    to count the occurrence frequencies of the words in a string's content
    :param word_list:
    :param content:
    :return:
    """
    occurrence = {}
    for i in word_list:
        occurrence[i] = content.count(i)
    return occurrence


def findSingleElement(candidate_iterable, search_str):
    """
    find the first element of the candidate iterable contains the search string
    :param candidate_iterable: list, tuple or other iterable with elements of string
    :param search_str: which string the element contains
    :return: the string element of the iterable contains the search string
    """
    for i in candidate_iterable:
        if search_str in i:
            return i
    return None


def doc2Docx(doc_path):
    """
    convert .doc file to .docx
    :param doc_path:
    :return:
    """
    # 打开.doc文件
    output_path = os.path.splitext(doc_path)[0] + '.docx'
    if os.path.exists(doc_path) and os.path.isfile(doc_path):
        with open(doc_path, 'rb') as file:
            document = Document(file)

            # 将.doc文件转换为.docx文件
        with open(output_path, 'wb') as file:
            document.save(file)
    else:
        print("please check the file")


def doc2DocxInFolder(path):
    """find all the .doc files recursively in a path, and convert them all to .docx files"""
    doc_list = findAllTypeFiles(path, 'doc')
    for doc in doc_list:
        doc2Docx(doc)
        os.remove(doc)


def openFileSafely(file_path):
    """
    when ever try opening a file, make sure it will be safely guarded even though not exists
    :param file_path: the path of the file to be opened
    :return:
    """
    if not os.path.exists(file_path):  # if not exists
        # create the directory
        dirs = os.path.dirname(file_path)
        if not dirs:  # local relative folder, result will be a void string ''
            dirs = '.'
        os.makedirs(dirs, exist_ok=True)
        # create the file
        fd = os.open(file_path, os.O_CREAT)
        # close the file handle
        os.close(fd)


def pics2Docx(pics_path, docx_path='pic.docx'):
    document = Document()
    for i in findFilesRecurisive(pics_path):
        document.add_picture(i, width=Cm(15), height=Cm(21))
    document.save(docx_path)


def nb2py(notebook):
    result = []
    cells = notebook['cells']

    for cell in cells:
        cell_type = cell['cell_type']

        if cell_type == 'markdown':
            result.append("%s'''\n%s\n'''" %
                          (header_comment, ''.join(cell['source'])))

        if cell_type == 'code':
            result.append("%s%s" % (header_comment, ''.join(cell['source'])))

    return '\n\n'.join(result)


def py2nb(py_str):
    # remove leading header comment
    if py_str.startswith(header_comment):
        py_str = py_str[len(header_comment):]

    cells = []
    chunks = py_str.split(header_comment)

    for chunk in chunks:
        if chunk and chunk.strip():
            chunk = '#' + chunk
            cell_type = 'code'
            if chunk.startswith("'''"):
                chunk = chunk.strip("'\n")
                cell_type = 'markdown'

            cell = {
                'cell_type': cell_type,
                'metadata': {},
                'source': chunk.splitlines(True),
            }

            if cell_type == 'code':
                cell.update({'outputs': [], 'execution_count': None})

            cells.append(cell)

    notebook = {
        'cells': cells,
        'metadata': {
            'anaconda-cloud': {},
            'kernelspec': {
                'display_name': 'Python 3',
                'language': 'python',
                'name': 'python3'},
            'language_info': {
                'codemirror_mode': {'name': 'ipython', 'version': 3},
                'file_extension': '.py',
                'mimetype': 'text/x-python',
                'name': 'python',
                'nbconvert_exporter': 'python',
                'pygments_lexer': 'ipython3',
                'version': '3.8.8'}},
        'nbformat': 4,
        'nbformat_minor': 1
    }

    return notebook


def convertInterPyNb(in_file, out_file):
    """
    convert a python notebook to a jupyter notebook or vice versa
    :param in_file:
    :param out_file:
    :return:
    """
    # make sure both type of files with the same stem name are not already in the same folder
    if in_file.rsplit('.', maxsplit=1)[0] == out_file.rsplit('.', maxsplit=1)[0]:
        return

    _, in_ext = path.splitext(in_file)
    _, out_ext = path.splitext(out_file)

    if in_ext == '.ipynb' and out_ext == '.py':
        with open(in_file, 'r') as f:
            notebook = json.load(f)
        py_str = nb2py(notebook)
        with open(out_file, 'w') as f:
            f.write(py_str)

    elif in_ext == '.py' and out_ext == '.ipynb':
        with open(in_file, 'r', encoding='utf8') as f:
            py_str = f.read()
        notebook = py2nb(py_str)
        with open(out_file, 'w', encoding='utf8') as f:
            json.dump(notebook, f, indent=2)

    else:
        raise (Exception('Extensions must be .ipynb and .py or vice versa'))


def batchConvertPy2nb(path, keep_source=True):
    """
    to batch convert all the python files in the path to jupyter notebook files
    :param path:
    :param keep_source: whether to keep the source python file
    :return:
    """
    pyfiles = findAllTypeFiles(path, extensions='py')
    for file in pyfiles:
        nbfile = file.rstrip('py') + 'ipynb'
        convertInterPyNb(file, nbfile)
        if not keep_source:
            os.remove(file)


def batchConvertNb2py(path, keep_source=True):
    """
    to batch convert all the jupyter notebook files in the path to python files
    :param path:
    :param keep_source: whether to keep the source notebook file
    :return:
    """
    nbfiles = findAllTypeFiles(path, extensions='ipynb')
    for file in nbfiles:
        pyfile = file.rstrip('ipynb') + 'py'
        convertInterPyNb(file, pyfile)
        if not keep_source:
            os.remove(file)


if __name__ == '__main__':
    # kw = 'Unable to download pixel_3a device skin; is your computer offline?'
    # pyfiles = findAllTypeFiles(r'C:\Users\ZXC\AppData\Local\Programs\Python\Python310\Lib\site-packages', '.py')
    # res = searchTextInFiles(kw, pyfiles)
    # print(res)

    # args = ArgsPass()
    # args.addArgs('account', str, 'simon')
    # account = args.readArgs('account')
    # print(account)
    ...
