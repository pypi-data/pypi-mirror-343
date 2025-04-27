# author : simon zhang
# date : 2022/05/11

import os, time
import re, random
from concurrent.futures import ThreadPoolExecutor
import requests
# from Crypto import Random
# from Crypto.Cipher import AES
from bs4 import BeautifulSoup
from docx import Document
from simonPackage.rw import findAllTypeFiles
import asyncio, aiofiles, aiohttp
from selenium.webdriver import Chrome
from selenium.webdriver.common.by import By

def myRequestGet(*args, **kwargs):
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:95.0) Gecko/20100101 Firefox/95.0'}
    proxy = None
    return requests.get(headers=headers, proxies=proxy, *args, **kwargs)

def myRequestPost(*args, **kwargs):
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:95.0) Gecko/20100101 Firefox/95.0'}
    proxy = None
    return requests.post(headers=headers, proxies=proxy, *args, **kwargs)


def getProxyList(proxy_web = 'https://www.kuaidaili.com/free/'):
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:95.0) Gecko/20100101 Firefox/95.0'}
    con = requests.get(proxy_web, headers=headers).content.decode(encoding='utf8')
    soup = BeautifulSoup(con, 'html.parser')
    test = soup.select('body div table tbody tr')
    proxy_list = []
    for i in test:
        ip = i.find('td', attrs={'data-title': "IP"}).text.strip()
        port = i.find('td', attrs={'data-title': "PORT"}).text.strip()
        tcp_type = i.find('td', attrs={'data-title': "类型"}).text.strip().lower()
        proxy = {tcp_type: ip + ':' + port}
        proxy_list.append(proxy)
    return proxy_list


def generateUrl(general_url, unique_url):
    """动态与页数拼接，生成具体的url
        para:
            general_url:str  通用一般url
            unique_url:int, str
       return
            url:str """
    url = f'{general_url}{unique_url}'
    return url


def getHtmlContent(url: str) -> str:
    """生成经过解码的包含正常中文文本字符串
        para:
           url:str  网址字符串
       return:
           html:str

           """
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:95.0) Gecko/20100101 Firefox/95.0'}
    html = myRequestGet(url).content.decode(encoding='utf8')
    print('getHtmlContent函数运行完成')
    return html


def postHtmlContent(url, *args, **kwargs):
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:95.0) Gecko/20100101 Firefox/95.0'}
    html = myRequestPost(url, *args, **kwargs).content.decode(encoding='utf8')
    return html



def getMainText(html):
    """得到页面的主要文字内容"""
    soup = BeautifulSoup(html, 'html.parser')
    title = soup.title.text
    main_text = title + '\n'
    for i in soup.find_all('p')[1:]:
        main_text += i.text
    return main_text


def getM3u8(url):
    """生成单个网页的m3u8网址和标题
    url:string
    return m3u8:string
           title: string
           """
    res = getHtmlContent(url)
    soup = BeautifulSoup(res, 'html.parser')
    pattern = r"(https.*?m3u8)"
    m3u8_list = re.findall(pattern, res)
    title = soup.head.title.text  # 通常可能的标题


    m3u8 = ''
    """特殊的标题方法"""
    match = re.search(r'正在播放\W(?P<name>.*?)-', res)
    if match:
        title = match.group('name')

    if len(m3u8_list):
        m3u8 = m3u8_list[0].replace('\\', '')

    else:
        print('no m3u8 url')


    # print(start, title)
    '''有可能第一次获取的不是真实的m3u8，打开后里面没有ts文件地址，而是包含了新的一个
    m3u8地址'''
    print(m3u8, title)
    return m3u8, title


def getMultiM3u8(number):
    """生成一定范围内页面的m3u8网址和标题"""
    unique_url = f'{number}-play.html?{number}-0-1'
    html = generateUrl(general_url='http://www.87el.xyz/AAyidong/AAAbf/', unique_url=unique_url)
    url, title = getM3u8(html)

    return url, title

def getTrueM3u8(m3u8):
    parts = m3u8.rpartition('/')
    true_m3u8 = parts[0] + '/hls/' + parts[2]
    return true_m3u8

def getListOfTs(m3u8, pattern_video=r'(.*?ts)'):
    """生成ts文件列表和key（如果key存在的话）
        para m3u8: string
       pattern_video:string  正则表达式规则
       return
            list_ts:list for ts in m3u8
            key:str   if key in encrypted file
        第一次获得的TS文件名可能只是相对名，需要拼接"""
    res = getHtmlContent(m3u8)
    # with open('aa.m3u8', 'wt') as f:
    #     f.write(res)
    m3u8dir = m3u8.rpartition('/')[0] + '/'
    raw_list = re.findall(pattern_video, res)
    list_ts = []
    for ts in raw_list:
        ts = ts.replace('\n', '').replace('\r', '')
        if ts.startswith('http'):
            # if ts.startswith(m3u8dir):
            pass

        else:
            list_ts.append(m3u8dir + ts)
    print(list_ts)
    key = None
    if 'key' in res:
        keylist = re.findall(r'"(https:.*?key.key)"', res)
        keyurl = keylist[0]
        key = myRequestGet(keyurl).content.decode(encoding='utf8')
        print('key is ', key)
    print(len(list_ts))
    print('getListOfTs函数运行完成')
    return list_ts, key


def get_Ts(ts):
    """得到单个的ts，写入ts文件中"""
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:95.0) Gecko/20100101 Firefox/95.0'}
    with open(ts, 'ab') as f:
        try:
            video_ts = myRequestGet(ts).content
            f.write(video_ts)
            print('reading ', ts)

        except:
            print('something is wrong with ', ts)


def getVideo(list_res, title):
    """读取列表中所有ts文件，并保存到指定文件中，文件名为title.mp4"""
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:95.0) Gecko/20100101 Firefox/95.0'}
    save_dir = f'./video_download/{title}'
    if not os.path.exists(save_dir):
        os.mkdir(save_dir)
    video_name = save_dir + '/' + str(title) + '.mp4'
    len_list = len(list_res)
    with open(video_name, 'wb') as f:
        for i in list_res:
            try:
                video_ts = myRequestGet(i).content
                f.write(video_ts)
                print('reading number', len_list, list_res.index(i), i)

            except:
                print('something is wrong with ', i)
                continue
            # print('written number', list_res.index(i), i)


# def decrypto(list_ts, key, title):
#     """对获取到的ts列表进行解密，返回真正的ts列表"""
#     numbers = len(list_ts)
#     for ts in list_ts:
#         res = myRequestGet(ts)
#         iv = Random.new().read(AES.block_size)
#         cryptor = AES.new(key.encode('utf-8'), AES.MODE_CBC, iv)
#         with open(f'{title}.mp4', 'ab') as f:
#             f.write(cryptor.decrypt(res.content))
#         print(f'finished number {numbers}/{str(list_ts.index(ts))}:  {ts}')

def joinTss(video_name, downloaded_ts_list):
    with open(video_name, 'ab') as vf:
        for ts in downloaded_ts_list:
            with open(ts, 'rb') as tsf:
                vf.write(tsf.read())
            os.remove(ts)

def get_Tss(ts, save_dir):
    ts_file = save_dir + ts.rpartition('/')[-1]
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:95.0) Gecko/20100101 Firefox/95.0'}
    with open(ts_file, 'ab') as f:
        try:
            video_ts = myRequestGet(ts).content
            f.write(video_ts)
            print('reading number', ts)

        except:
            print('something is wrong with ', ts)

    print(f"{ts}下载完成")

def download_Videos(url):
    m3u8, title = getM3u8(url)
    list_ts = getListOfTs(m3u8)[0]
        # if not the true m3u8 file, wouldn't get any ts
    if not len(list_ts):
        m3u8 = getTrueM3u8(m3u8)
        list_ts = getListOfTs(m3u8)[0]
    print(list_ts)
    save_dir = f'./video_download/{title}/'
    if not os.path.exists(save_dir):
        os.mkdir(save_dir)
    video_name = save_dir + str(title) + '.mp4'
    with ThreadPoolExecutor(100) as t:
        for ts in list_ts:
            t.submit(get_Tss, ts=ts, save_dir=save_dir)

    # 关闭线程池。这会等待所有正在进行的任务完成。
    t.shutdown(wait=True)
    downloaded_ts_list = findAllTypeFiles(save_dir, extentions='ts')
    downloaded_sorted_ts_list = []
    for ts in list_ts:
        name = ts.rpartition('/')[-1]
        for downloaded_ts in downloaded_ts_list:
            if name in downloaded_ts:
                downloaded_sorted_ts_list.append(downloaded_ts)

    print(downloaded_ts_list)
    joinTss(video_name, downloaded_sorted_ts_list)
    print('all done', title)


def download_Videos_Single_Thread(url):
    m3u8, title = getM3u8(url)
    m3u8 = getTrueM3u8(m3u8)
    list_ts = getListOfTs(m3u8)[0]
    print(list_ts)
    save_dir = f'./video_download/{title}/'
    if not os.path.exists(save_dir):
        os.mkdir(save_dir)
    video_name = save_dir + str(title) + '.mp4'
    with open(video_name, 'ab') as f:
        count = 1
        total = len(list_ts)
        for ts in list_ts:
            try:
                f.write(myRequestGet(ts).content)
                print(f'{total}/{count}: {ts}  done')
                count += 1
            except:
                print(f'something is wrong with {ts}')
                continue

    print('all done')

async def download_ts(ts, session):
    name = ts.rpartition('/')[-1]
    async with session.get(ts) as resp:
        async with aiofiles.open(f'./videodownload/{name}', 'wb') as f:
            await f.write(await resp.content.read())

async def aioDownload(list_ts):
    tasks = []
    async with aiohttp.ClientSession() as session:
        async for ts in list_ts:
            task = asyncio.create_task(download_ts(ts, session))
            tasks.append(task)

def getElementText(markup, tag_name):
    soup = BeautifulSoup(markup, 'html.parser')
    # print(soup.find(tag_name).text)
    return soup.find(tag_name).text

def getJsonData(str):
    jsonData = {}
    list_kv = str.split('\n')
    for kv in list_kv:
        k= kv.split(':')[0]
        v= kv.split(':')[1]
        jsonData[k] = v
    return jsonData


def saveDocument(url, path='.'):
    """
    save the url's content to a docx file
    :param url:
    :return:
    """
    if path == '.':
        path = os.path.abspath(os.curdir)
        
    def setText(run, text, size=177800):
        run.font.name = '仿宋_GB2312'
        run.font.size = size
        run.text = text
        
    res = getHtmlContent(url)
    # content = getMainText(res)
    soup = BeautifulSoup(res, 'html.parser')
    title = soup.h1.text if soup.h1 else soup.title.text
    title = title.strip()
    doc = Document()
    doc.add_heading(title)
    for para in soup.find_all('p'):
        para = doc.add_paragraph('\u3000' * 2 + para.text)
        for run in para.runs:
            # print(run.text)
            setText(run, run.text)
    file_path = f'{path}/{title}.docx'
    doc.save(file_path)
    return file_path