import difflib
import os
import re
import traceback
import base64
from .read_pdf import read_pdf
from PIL import Image
from io import BytesIO
import glob
class Files:
    def __init__(self, computer):
        self.computer = computer

    def upload(self, file_path):
        return self.computer.ai.upload(file_path)

    def _read_file(self, path, pdf_image_start=0, pdf_image_end=1):
        # Get file extension
        _, ext = os.path.splitext(path)
        ext = ext.lower()

        # Handle PDFs
        if ext == '.pdf':
            return read_pdf(path, pdf_image_start=pdf_image_start, pdf_image_end=pdf_image_end)

        # Handle DOCX files
        elif ext == '.docx':
            from docx import Document

            doc = Document(path)
            output = ""

            # Loop through paragraphs, print index and text
            for i, paragraph in enumerate(doc.paragraphs):
                output += f'Paragraph [{i}]: "{paragraph.text}"\n'

            # Directly show table contents and index if present
            for tbl_idx, table in enumerate(doc.tables):
                print(f'---Table [{tbl_idx}]---')
                for row_idx, row in enumerate(table.rows):
                    row_data = [f'Cell[{cell_idx}]: "{cell.text}"' for cell_idx, cell in enumerate(row.cells)]
                    output += f"Row [{row_idx}]" + ' | '.join(row_data) + "\n"

            print("\n\nIf you want to edit this, please use python-docx to manipulate this content programmatically.")

            return {
                "text": output
            }

        # Handle images
        elif ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
            with open(path, 'rb') as file:
                image = Image.open(file)
                width, height = image.size
                
                # Scale down if needed
                if width > 1000 or height > 1000:
                    scale = min(1000/width, 1000/height)
                    new_width = int(width * scale)
                    new_height = int(height * scale)
                    image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
                
                # Convert to base64
                img_buffer = BytesIO()
                image.save(img_buffer, format=image.format or 'PNG')
                base64_img = base64.b64encode(img_buffer.getvalue()).decode()
                # Get format from extension, removing the dot and defaulting to png
                img_format = ext[1:] if ext else 'png'
                # Construct the data url based on the image extension
                data_url = f"data:image/{img_format};base64,{base64_img}"

            return {
                "base64": [f"<lt_base64>{data_url}</lt_base64>"]
            }

        # Handle audio files
        elif ext in ['.wav', '.mp3', '.m4a']:
            return {"text": self.computer.audio.transcribe(path)}

        # Handle text files
        else:
            with open(path, 'r') as file:
                data = file.read()
            return {
                "text": data
            }

    def read(self, path, extensions=None, pdf_image_start=0, pdf_image_end=1):
        if path is None or path == ".":
            path = os.getcwd()

        # Make it so path handles wildcards like /Users/killianlucas/Desktop/Screenshot*2025-02-07*at*1.20.17*PM.png
        if "*" in path:
            matches = glob.glob(path)
            if not matches:
                raise ValueError(f"No files found matching pattern: {path}")
            path = matches[0]

        if extensions is None:
            extensions = (
                ".txt", ".md", ".py", ".js", ".html", ".css", ".json", ".yaml", ".yml",
                ".ini", ".cfg", ".conf", ".toml", ".xml", ".csv", ".tsv",
                ".rst", ".adoc", ".tex", ".sh", ".bash", ".zsh", ".fish",
                ".sql", ".php", ".rb", ".pl", ".pm", ".go", ".java", ".kt", ".scala",
                ".c", ".cpp", ".h", ".hpp", ".cs", ".ts", ".swift", ".m", ".mm",
                ".r", ".lua", ".tcl", ".f", ".f90", ".hs", ".erl", ".ex", ".exs",
                ".clj", ".groovy", ".ps1", ".vbs", ".bat", ".cmd"
            )

        if os.path.isfile(path):
            return self._read_file(path, pdf_image_start=pdf_image_start, pdf_image_end=pdf_image_end)

        elif os.path.isdir(path):
            file_list = []
            files_that_dont_match_extension = []
            for root, _, files in os.walk(path):
                for file in files:
                    file_path = os.path.join(root, file)
                    if not file.endswith(extensions):
                        files_that_dont_match_extension.append(file_path)
                        print(f"Skipping file: {file_path} (does not match specified extensions)")
                        print("To include this file type, pass in a tuple of extensions to the 'extensions' parameter, e.g., extensions=('.py', '.txt', '.md')")
                        continue
                    try:
                        file_text = self._read_file(file_path)["text"]
                    except Exception as e:
                        print(
                            f"Failed to read file at path: {file_path}. Reason: {str(e)}"
                        )
                        continue
                    file_list.append({"path": file_path, "text": file_text})
            if file_list == []:
                raise Exception(f"No files with the extensions: {extensions} found in the path: {path}. Instead, these files were found there: {files_that_dont_match_extension}")
            return file_list  # "\n".join([f["path"] + "\n" + f["text"] for f in file_list])
        else:
            raise ValueError(f"Path '{path}' is neither a file nor a directory.")

    def search(self, query, path=None, height=3):
        """
        Search the filesystem for the given query.
        """
        if path is None or path == ".":
            path = os.getcwd()

        if os.path.isfile(path):
            try:
                with open(path, 'r', encoding='utf-8') as file:
                    lines = file.readlines()
            except UnicodeDecodeError:
                # If UTF-8 fails, try with a more permissive encoding
                with open(path, 'r', encoding='latin-1') as file:
                    lines = file.readlines()
            except Exception as e:
                raise Exception(f"Could not read path {path}: {str(e)}")
            matches = []
            for i, line in enumerate(lines):
                if query in line:
                    start = max(0, i - height)
                    end = min(len(lines), i + height + 1)
                    context = ''.join(lines[start:end])
                    if len(context) > 300:
                        index = context.find(query)
                        start = max(0, index - 150)
                        end = min(len(context), index + 150)
                        context = context[start:end]
                        if start > 0:
                            context = '...' + context
                        if end < len(context):
                            context = context + '...'
                    matches.append({
                        'path': path,
                        'line_number': i + 1,
                        'line': line.strip(),
                        'context': context
                    })
            return matches
        elif os.path.isdir(path):
            matches = []
            for root, _, files in os.walk(path):
                for file in files:
                    file_path = os.path.join(root, file)
                    matches.extend(self.search(query, file_path, height))
            if not matches:
                close_matches = get_close_matches_in_text(query, path)
                if close_matches:
                    raise ValueError(f"No exact matches found. Did you mean one of these? {close_matches}")
                else:
                    # print("Using flexible AI search, which might be slow...")
                    search_for = query
                    return self.files.query(query=f"Can you find all instances of something very close to this text:\n\n'''{search_for}'''\n\nWhere is that in the files? Be SPECIFIC as to the exact line(s) where it appears, and in what file(s).", path=path)
            else:
                return matches
        else:
            raise ValueError(f"Path '{path}' is neither a file nor a directory.")

    def query(self, query, paths=None, extensions=None, model_size="tiny"):
        """
        Ask a question about the files in the given path. Doesn't view images in PDFs.
        """
        if isinstance(paths, str):
            if "," in paths:
                paths = [p.strip() for p in paths.split(",")]
            else:
                paths = [paths]
        elif isinstance(paths, list):
            paths = paths
        else:
            raise ValueError("Path must be a string or list of strings")
            
        # Build up text from all paths
        text = ""
        for p in paths:
            result = self.read(p, extensions, pdf_image_start=0, pdf_image_end=0)
            
            # Add any text content
            if "text" in result:
                text += f"\n==The following is the text of {p}==\n{result['text']}\n\n"
                
            # Process any base64 images
            if "base64" in result:
                for b64_str in result["base64"]:
                    # Remove <lt_base64> tags from base64 string
                    b64_str = b64_str.replace("<lt_base64>", "").replace("</lt_base64>", "")
                    image_analysis = self.computer.ai.chat(query, base64_image=b64_str)
                    text += f"\n==The following is an analysis of the image at {p}==\n{image_analysis}\n\n"

            # If we only got a single image from a single path, return analysis directly
            if (len(paths) == 1 and 
                not result.get("text") and 
                result.get("base64") and 
                len(result["base64"]) == 1):
                return image_analysis.strip()
                    
        answer = self.computer.ai.query(text, query, model_size=model_size)
        return answer.strip()

    def edit(self, path, original_text, replacement_text, force=False):
        """
        Edits a file on the filesystem, replacing the original text with the replacement text.
        """
        with open(path, "r") as file:
            filedata = file.read()

        if original_text not in filedata:
            matches = get_close_matches_in_text(original_text, filedata)
            if matches:
                suggestions = ", ".join(matches)
                raise ValueError(
                    f"Original text not found. Did you mean one of these? {suggestions}"
                )

        if force == False:
            
            for attempt in range(10):
                # Perform the replacement
                new_filedata = filedata.replace(original_text, replacement_text)

                syntax_errors = check_syntax(new_filedata)

                # Get context around the edit
                old_lines = filedata.splitlines()
                new_lines = new_filedata.splitlines()

                # Find the start and end of the edit
                edit_start = next(i for i, (old, new) in enumerate(zip(old_lines, new_lines)) if old != new)
                edit_end = next(i for i, (old, new) in enumerate(zip(old_lines[::-1], new_lines[::-1])) if old != new)
                edit_end = len(old_lines) - edit_end

                # Get 50 lines of context on either side
                context_start = max(0, edit_start - 50)
                context_end = min(len(new_lines), edit_end + 50)

                # Generate the diff
                diff = []
                diff.append("...")
                diff.append("")

                # Add 50 lines of context before the edit
                for i in range(context_start, edit_start):
                    diff.append(new_lines[i])

                diff.append("--- OLD CODE START ---")
                # Add all the old lines
                for i in range(edit_start, edit_end):
                    if i < len(old_lines):
                        diff.append(old_lines[i])
                diff.append("--- OLD CODE END ---")
                diff.append("--- NEW CODE START ---")
                # Add all the new lines
                for i in range(edit_start, edit_end):
                    if i < len(new_lines):
                        diff.append(new_lines[i])
                diff.append("--- NEW CODE END ---")

                # Add 50 lines of context after the edit
                for i in range(edit_end, context_end):
                    diff.append(new_lines[i])

                context = "\n".join(diff)

                # Prepare LLM query
                editor_system_message = "You are an expert Python programmer. Please review the following code edit and focus solely on ensuring the code is placed correctly within the file. The content of the code is correct, but its placement or indentation might be unintentional."
                editor_user_message = f"Please review this code edit, focusing only on whether its placement seems intentional or unintentional. Here's the context around the edit, including 50 lines before and after:\n\n{context}\n\n"

                if syntax_errors:
                    editor_user_message += f"This edit resulted in syntax errors. Here are the errors:\n\n{syntax_errors}\n\n"

                editor_user_message += "If the edit's placement appears to be intentional and correct, simply respond with 'Approved'. If you believe the placement is unintentional or incorrect, please provide suggestions to correct it wrapped in <find>old code</find> and <replace>new code</replace> tags. Focus only on placement and indentation, not on the content or style of the code."

                print("---EDIT QUERY---")
                print(editor_user_message)
                print("---END OF EDIT QUERY---")

                messages = [
                    {
                        "role": "system",
                        "type": "message",
                        "content": editor_system_message,
                    },
                    {"role": "user", "type": "message", "content": editor_user_message},
                ]
                
                print("\n--- APPLYING EDIT ---\n")
                report = self.computer.ai.chat(messages=messages, display=True)
                print("\n--- DONE APPLYING EDIT ---\n")

                # Check if the LLM approved the edit
                if "<find>" not in report and "<replace>" not in report:
                    break  # Exit the loop if the LLM approved the edit

                # Extract find and replace from the LLM response
                find_pattern = re.compile(r'<find>(.*?)</find>', re.DOTALL)
                replace_pattern = re.compile(r'<replace>(.*?)</replace>', re.DOTALL)
                
                find_match = find_pattern.search(report)
                replace_match = replace_pattern.search(report)
                
                if find_match and replace_match:
                    original_text = find_match.group(1)
                    replacement_text = replace_match.group(1)
                else:
                    print("LLM did not provide a valid fix or approval!")
                    continue
            else:
                raise ValueError(f"AI seemed to never approved this edit! Last AI message: {report}\n\nIf this seems to be an error, pass in force=True to computer.files.edit.")

        elif force == True:
            # Perform the replacement
            new_filedata = filedata.replace(original_text, replacement_text)

        # Apply the final approved edit
        with open(path, "w") as file:
            file.write(new_filedata)

    def map(self, path=None, extensions=None, codebase=None):
        if extensions == None:
            print("Defaulting extensions to ('.py'), pass in a tuple with other extensions to see more, but this should be sufficient.")
            extensions = ('.py')

        if path is None or path == ".":
            path = os.getcwd()
        
        if not os.path.isdir(path):
            raise ValueError(f"The provided path '{path}' is not a directory.")

        if codebase == None:
            codebase = is_codebase(path)

        def generate_tree(directory, prefix="", extensions=None):
            tree = ""
            contents = sorted(os.listdir(directory))
            for i, item in enumerate(contents):
                path_item = os.path.join(directory, item)
                if os.path.isdir(path_item):
                    sub_tree = generate_tree(path_item, prefix + "│   ", extensions)
                    if (
                        sub_tree.strip()
                    ):  # Only add the folder if it contains files with the specified extensions
                        tree += f"{prefix}├── {item}/\n"
                        tree += sub_tree
                else:
                    if extensions is None or item.endswith(extensions):
                        tree += f"{prefix}├── {item}\n"
                        if codebase and (
                            extensions is None or item.endswith(extensions)
                        ):
                            tree += generate_codebase_info(path_item, prefix + "│   ")
            return tree

        def generate_codebase_info(file_path, prefix=""):
            codebase_info = ""
            try:
                with open(file_path, "r", errors='ignore') as file:
                    lines = file.readlines()
                    for line in lines:
                        stripped_line = line.strip()
                        if stripped_line.startswith("class ") or stripped_line.startswith(
                            "def "
                        ):
                            codebase_info += f"{prefix}│   {stripped_line}\n"
            except UnicodeDecodeError:
                codebase_info += f"{prefix}│   [Error: Unable to read file]\n"
            except Exception as e:
                codebase_info += f"{prefix}│   [Error: {str(e)}]\n"
            return codebase_info

        return generate_tree(path, extensions=extensions)

def get_close_matches_in_text(original_text, filedata, n=3):
    """
    Returns the closest matches to the original text in the content of the file.
    """
    words = filedata.split()
    original_words = original_text.split()
    len_original = len(original_words)

    matches = []
    for i in range(len(words) - len_original + 1):
        phrase = " ".join(words[i : i + len_original])
        similarity = difflib.SequenceMatcher(None, original_text, phrase).ratio()
        matches.append((similarity, phrase))

    matches.sort(reverse=True)
    return [match[1] for match in matches[:n]]

def is_codebase(path):
    """
    Determine if the given path is part of a codebase (i.e., contains a .git directory).
    
    Args:
    path (str): The path to check. Can be a file or directory.
    
    Returns:
    bool: True if the path is part of a codebase, False otherwise.
    """
    # If it's a file, return False immediately
    if os.path.isfile(path):
        return False
    
    # Traverse up the directory tree
    while path != os.path.dirname(path):  # Stop at the root directory
        if os.path.exists(os.path.join(path, ".git")):
            return True
        path = os.path.dirname(path)
    
    return False

def check_syntax(code_string):
    try:
        compile(code_string, '<string>', 'exec')
        return None
    except SyntaxError:
        return traceback.format_exc()
