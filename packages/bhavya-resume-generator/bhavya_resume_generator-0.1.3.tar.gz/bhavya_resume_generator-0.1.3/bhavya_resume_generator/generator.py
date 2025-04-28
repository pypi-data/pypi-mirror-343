from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.enum.table import WD_ALIGN_VERTICAL

class ResumeGenerator:
    def __init__(self, add_h1b_header=False):
        self.doc = Document()
        self.section = self.doc.sections[0]
        self._setup_document()
        self.add_h1b_header = add_h1b_header
        if add_h1b_header:
            self._add_h1b_header()

    def _setup_document(self):
        """Set up document margins and borders"""
        self._set_margins()
        self._add_page_border()

    def _set_margins(self):
        """Set document margins"""
        self.section.top_margin = Inches(0.2)
        self.section.bottom_margin = Inches(0.2)
        self.section.left_margin = Inches(0.2)
        self.section.right_margin = Inches(0.2)
        self.section.header_distance = Inches(0.2)

    def _add_page_border(self):
        """Add border to the page"""
        sectPr = self.section._sectPr
        pgBorders = parse_xml(r'<w:pgBorders %s>'
                          '<w:top w:val="single" w:sz="12" w:space="4" w:color="000000"/>'
                          '<w:left w:val="single" w:sz="12" w:space="4" w:color="000000"/>'
                          '<w:bottom w:val="single" w:sz="12" w:space="4" w:color="000000"/>'
                          '<w:right w:val="single" w:sz="12" w:space="4" w:color="000000"/>'
                          '</w:pgBorders>' % nsdecls('w'))
        sectPr.append(pgBorders)

    def _add_h1b_header(self):
        """Add H1B header to the document"""
        header = self.section.header
        if not header.paragraphs:
            paragraph = header.add_paragraph()
        else:
            paragraph = header.paragraphs[0]
        paragraph.text = "H1B"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.runs[0]
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0

    def add_candidate_details(self, name, email, phone, location, linkedin=None):
        """Add candidate details to the resume"""
        # Name
        name_para = self._add_with_padding()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(name.upper())
        name_run.font.size = Pt(18)
        name_run.font.bold = True
        name_run.font.name = 'Times New Roman'
        name_para.paragraph_format.space_after = Pt(0)

        # Contact - Email and Phone on one line
        contact_para = self._add_with_padding()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        email_run = contact_para.add_run(email)
        email_run.font.size = Pt(11)
        email_run.font.name = 'Times New Roman'
        contact_para.add_run(' | ')
        phone_run = contact_para.add_run(phone)
        phone_run.font.size = Pt(11)
        phone_run.font.name = 'Times New Roman'
        contact_para.paragraph_format.space_before = Pt(0)
        contact_para.paragraph_format.space_after = Pt(2)
        contact_para.paragraph_format.line_spacing = 1.0

        # Location
        if location:
            loc_para = self._add_with_padding()
            loc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            loc_run = loc_para.add_run(location)
            loc_run.font.size = Pt(11)
            loc_run.font.name = 'Times New Roman'
            loc_para.paragraph_format.space_before = Pt(0)
            loc_para.paragraph_format.space_after = Pt(2)
            loc_para.paragraph_format.line_spacing = 1.0

        # LinkedIn
        if linkedin:
            linkedin_para = self._add_with_padding()
            linkedin_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            linkedin_run = linkedin_para.add_run(linkedin)
            linkedin_run.font.size = Pt(11)
            linkedin_run.font.name = 'Times New Roman'
            linkedin_para.paragraph_format.space_before = Pt(0)
            linkedin_para.paragraph_format.space_after = Pt(2)
            linkedin_para.paragraph_format.line_spacing = 1.0

    def add_summary(self, summary_text):
        """Add professional summary section"""
        self._add_section_title('Professional Summary')
        self._add_bullets(summary_text)

    def add_education(self, degree, university, graduation_year=None):
        """Add education section"""
        self._add_section_title('Education')
        para = self._add_with_padding()
        run = para.add_run('Education Degree: ')
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        para.add_run(degree)
        
        if university:
            para.add_run('\n')
            run2 = para.add_run('University Name: ')
            run2.font.bold = True
            run2.font.size = Pt(11)
            run2.font.name = 'Times New Roman'
            para.add_run(university)
            
        if graduation_year:
            para.add_run('\n')
            run3 = para.add_run('Graduation Year: ')
            run3.font.bold = True
            run3.font.size = Pt(11)
            run3.font.name = 'Times New Roman'
            para.add_run(graduation_year)

    def add_skills(self, skills_dict):
        """Add skills section as a table"""
        self._add_section_title('Technical Skills')
        table = self.doc.add_table(rows=len(skills_dict), cols=2)
        table.allow_autofit = False
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set table width
        table_width = Inches(7.7)
        col_width = table_width / 2
        
        for i, (category, details) in enumerate(skills_dict.items()):
            row = table.rows[i]
            for cell in row.cells:
                cell.width = col_width
                
            cat_cell = row.cells[0]
            det_cell = row.cells[1]
            
            cat_para = cat_cell.paragraphs[0]
            det_para = det_cell.paragraphs[0]
            
            cat_run = cat_para.add_run(category.strip())
            det_run = det_para.add_run(details.strip())
            
            det_run.font.size = Pt(11)
            det_run.font.name = 'Times New Roman'
            
            self._set_cell_margins(cat_cell, top=10, start=288, bottom=10, end=288)
            self._set_cell_margins(det_cell, top=10, start=288, bottom=10, end=288)
            
            cat_para.paragraph_format.space_after = Pt(0)
            det_para.paragraph_format.space_after = Pt(0)
            
        self._add_table_borders(table)

    def add_experience(self, experience_entries):
        """Add professional experience section"""
        self._add_section_title('Professional Experience')
        
        for entry in experience_entries:
            # Company and Location
            para = self._add_with_padding()
            tab_stops = para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(6.5))
            
            company_loc = f"{entry.get('company', '')} | {entry.get('location', '')}"
            run = para.add_run(company_loc)
            run.font.bold = True
            run.font.underline = True
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            
            para.add_run("\t")
            dur_run = para.add_run(entry.get('duration', ''))
            dur_run.font.size = Pt(11)
            dur_run.font.name = 'Times New Roman'
            para.paragraph_format.space_after = Pt(0)
            
            # Position
            pos_para = self._add_with_padding()
            pos_run = pos_para.add_run(entry.get('position', ''))
            pos_run.font.italic = True
            pos_run.font.size = Pt(11)
            pos_run.font.name = 'Times New Roman'
            pos_para.paragraph_format.space_before = Pt(0)
            pos_para.paragraph_format.space_after = Pt(0)
            
            # Description bullets
            for idx, desc in enumerate(entry.get('description', [])):
                bullet_para = self._add_with_padding()
                bullet_para.paragraph_format.left_indent = Inches(0.5)
                bullet_para.paragraph_format.first_line_indent = Inches(-0.1)
                bullet_para.paragraph_format.space_before = Pt(0) if idx == 0 else Pt(2)
                bullet_para.paragraph_format.space_after = Pt(2)
                bullet_para.paragraph_format.line_spacing = 1.0
                run = bullet_para.add_run('• ' + desc)
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
            
            # Skills
            if entry.get('skills'):
                skills_para = self._add_with_padding()
                skills_run = skills_para.add_run('Skills: ' + entry['skills'])
                skills_run.font.italic = True
                skills_run.font.size = Pt(10)
                skills_run.font.name = 'Times New Roman'
                skills_run.font.color.rgb = RGBColor(120, 120, 120)
                skills_para.paragraph_format.space_before = Pt(0)
                skills_para.paragraph_format.space_after = Pt(8)

    def generate(self, output_file):
        """Generate the resume document"""
        self.doc.save(output_file)
        return output_file

    def _add_with_padding(self, paragraph=None):
        """Helper to add content with padding"""
        para = self.doc.add_paragraph() if paragraph is None else self.doc.add_paragraph(paragraph)
        para.paragraph_format.left_indent = Inches(0.2)
        para.paragraph_format.right_indent = Inches(0.2)
        return para

    def _add_section_title(self, title):
        """Add a section title"""
        para = self._add_with_padding()
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.line_spacing = 1.0
        run = para.add_run(title.upper())
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run.underline = True
        para.paragraph_format.keep_with_next = True
        return para

    def _add_bullets(self, text):
        """Add bullet points"""
        for line in text.split('\n'):
            if line.strip().startswith('-') or line.strip().startswith('•') or line.strip().startswith('.'):
                bullet_para = self._add_with_padding()
                bullet_para.paragraph_format.left_indent = Inches(0.5)
                bullet_para.paragraph_format.first_line_indent = Inches(-0.1)
                bullet_para.paragraph_format.space_before = Pt(0)
                bullet_para.paragraph_format.space_after = Pt(2)
                bullet_para.paragraph_format.line_spacing = 1.0
                clean_text = line.lstrip('-•.').strip()
                run = bullet_para.add_run('• ' + clean_text)
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
            elif line.strip():
                para = self._add_with_padding(line.strip())
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.line_spacing = 1.0
                run = para.runs[0]
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'

    def _set_cell_margins(self, cell, top=144, start=144, bottom=144, end=144):
        """Set cell margins"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = tcPr.find(qn('w:tcMar'))
        if tcMar is None:
            tcMar = OxmlElement('w:tcMar')
            tcPr.append(tcMar)
        for margin, value in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
            node = tcMar.find(qn(f'w:{margin}'))
            if node is None:
                node = OxmlElement(f'w:{margin}')
                tcMar.append(node)
            node.set(qn('w:w'), str(value))
            node.set(qn('w:type'), 'dxa')

    def _add_table_borders(self, table):
        """Add borders to a table"""
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.append(tblPr)
        borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            borders.append(border)
        tblPr.append(borders) 