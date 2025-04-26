from __future__ import annotations

from io import BytesIO
from typing import TYPE_CHECKING

import docx
from django.conf import settings
from django.utils.timezone import now

from pybmds.datasets.transforms.polyk import PolyKAdjustment
from pybmds.datasets.transforms.rao_scott import RaoScott
from pybmds.reporting.styling import Report, write_setting_p
from pybmds.utils import get_version

from ... import __version__ as bmds_ui_version
from ...common.docx import add_url_hyperlink
from ...common.utils import to_timestamp
from ..utils import get_citation

if TYPE_CHECKING:
    from ..models import Analysis

ANALYSIS_URL = "Analysis URL: "


def write_version_p(report: Report, bmds_ui: str, pybmds: str, bmdscore):
    version_label = "BMDS Desktop Version: " if settings.IS_DESKTOP else "BMDS Online Version: "
    version_str = f"{bmds_ui} (pybmds {pybmds}; bmdscore {bmdscore})"
    write_setting_p(report, version_label, version_str)


def build_docx(
    analysis: Analysis,
    uri: str,
    dataset_format_long: bool = True,
    all_models: bool = False,
    bmd_cdf_table: bool = False,
) -> BytesIO:
    """Generate a Microsoft Word binary file for an analysis

    Args:
        analysis (Analysis): An Analysis object
        uri (str): The root URI for this site, eg: "https://example.com"
        dataset_format_long (bool, default True): long or wide dataset table format
        all_models (bool, default False):  Show all models, not just selected
        bmd_cdf_table (bool, default False): Export BMD CDF table

    Returns:
        BytesIO: A word document byte stream
    """
    f = BytesIO()
    report = Report.build_default()

    report.document.add_heading(analysis.name(), 1)

    description = analysis.inputs.get("analysis_description")
    if description:
        report.document.add_paragraph(description)

    write_setting_p(report, "Report Generated: ", to_timestamp(now()))

    if not settings.IS_DESKTOP:
        p = report.document.add_paragraph()
        p.add_run(ANALYSIS_URL).bold = True
        uri += analysis.get_absolute_url()
        add_url_hyperlink(p, uri, "View")

    write_version_p(
        report,
        analysis.outputs["bmds_ui_version"],
        analysis.outputs["bmds_python_version"]["python"],
        analysis.outputs["bmds_python_version"]["dll"],
    )

    if not analysis.is_finished:
        report.document.add_paragraph("Execution is incomplete; no report could be generated")
    elif analysis.has_errors:
        report.document.add_paragraph("Execution generated errors; no report can be generated")
    else:
        batch = analysis.to_batch()
        batch.to_docx(
            report=report,
            header_level=1,
            citation=False,
            dataset_format_long=dataset_format_long,
            all_models=all_models,
            bmd_cdf_table=bmd_cdf_table,
            session_inputs_table=True,
        )

    write_citation(report, 1)

    report.document.save(f)
    return f


def write_citation(report: Report, header_level: int):
    styles = report.styles
    header_style = styles.get_header_style(header_level)
    report.document.add_paragraph("Recommended citation", header_style)
    report.document.add_paragraph(
        "Please adapt as appropriate; the citations below capture the package version and "
        "timestamps for easier reproducibility of the analysis."
    )
    report.document.add_paragraph(get_citation(), styles.fixed_width)


def add_update_url(analysis: Analysis, data: BytesIO, uri: str) -> BytesIO:
    """Add an update URL to an existing BMDS report

    Args:
        analysis (Analysis): An Analysis object
        data (BytesIO):  A word document byte stream
        uri (str): The root URI for this site, eg: "https://example.com"

    Returns:
        BytesIO: A word document byte stream
    """
    document = docx.Document(data)
    for p in document.paragraphs:
        if p.text.startswith(ANALYSIS_URL):
            p.add_run(" / ")
            uri += analysis.get_edit_url()
            add_url_hyperlink(p, uri, "Update")
            break

    f = BytesIO()
    document.save(f)
    return f


def write_current_version_p(report):
    versions = get_version()
    write_version_p(report, bmds_ui_version, versions.python, versions.dll)


def build_polyk_docx(analysis: PolyKAdjustment) -> BytesIO:
    report = Report.build_default()

    report.document.add_heading("Poly K Adjustment", 1)
    write_setting_p(report, "Report generated: ", to_timestamp(now()))
    write_current_version_p(report)
    analysis.to_docx(report=report, show_title=False)

    f = BytesIO()
    report.document.save(f)
    return f


def build_raoscott_docx(analysis: RaoScott) -> BytesIO:
    report = Report.build_default()

    report.document.add_heading("Rao-Scott Transformation", 1)
    write_setting_p(report, "Report generated: ", to_timestamp(now()))
    write_current_version_p(report)
    analysis.to_docx(report=report, show_title=False)

    f = BytesIO()
    report.document.save(f)
    return f
